#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Convierte los .md del proyecto TYM Music a .docx nativos para Word/Drive."""
import os, re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

AZUL = RGBColor(0x1F, 0x4E, 0x79)
AZUL2 = RGBColor(0x2E, 0x75, 0xB6)
BLANCO = RGBColor(0xFF, 0xFF, 0xFF)
GRIS = RGBColor(0x55, 0x55, 0x55)

def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear"); sh.set(qn("w:color"), "auto"); sh.set(qn("w:fill"), hexcolor)
    tcPr.append(sh)

def add_runs(par, text):
    """Procesa **bold**, `code` e inline."""
    text = re.sub(r"\[\[([^\]]+)\]\]", r"\1", text)          # [[links]]
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)      # [txt](url)
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text)
    for p in parts:
        if not p:
            continue
        if p.startswith("**") and p.endswith("**"):
            r = par.add_run(p[2:-2]); r.bold = True
        elif p.startswith("`") and p.endswith("`"):
            r = par.add_run(p[1:-1]); r.font.name = "Consolas"; r.font.size = Pt(9.5)
        else:
            par.add_run(p)

def split_row(line):
    cells = line.strip().strip("|").split("|")
    return [c.strip() for c in cells]

def is_sep(line):
    return bool(re.match(r"^\s*\|?[\s:\-|]+\|?\s*$", line)) and "-" in line

def setup(doc, title):
    n = doc.styles["Normal"]; n.font.name = "Calibri"; n.font.size = Pt(10.5)
    doc.styles["Heading 1"].font.color.rgb = AZUL
    doc.styles["Heading 1"].font.size = Pt(15)
    doc.styles["Heading 2"].font.color.rgb = AZUL2
    doc.styles["Heading 2"].font.size = Pt(12.5)
    try:
        doc.styles["Heading 3"].font.color.rgb = AZUL2
        doc.styles["Heading 3"].font.size = Pt(11)
    except KeyError:
        pass

def make_table(doc, header, rows):
    t = doc.add_table(rows=1, cols=len(header))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hc = t.rows[0].cells
    for i, h in enumerate(header):
        shade(hc[i], "1F4E79")
        pp = hc[i].paragraphs[0]; pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rr = pp.add_run(re.sub(r"\*\*", "", h)); rr.bold = True
        rr.font.color.rgb = BLANCO; rr.font.size = Pt(9.5)
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for i in range(len(header)):
            val = row[i] if i < len(row) else ""
            if ri % 2 == 1:
                shade(cells[i], "EAF1F8")
            pp = cells[i].paragraphs[0]
            add_runs(pp, val)
            for run in pp.runs:
                run.font.size = Pt(9.5)
            if i == 0:
                for run in pp.runs:
                    run.bold = True
    doc.add_paragraph()

def convert(src, dest, title, subtitle):
    with open(src, encoding="utf-8") as f:
        lines = f.read().split("\n")
    doc = Document()
    setup(doc, title)
    # Portada simple
    for _ in range(2): doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("TYM MUSIC"); r.bold = True; r.font.size = Pt(34); r.font.color.rgb = AZUL
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title); r.font.size = Pt(15); r.font.color.rgb = AZUL2
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(subtitle); r.italic = True; r.font.size = Pt(10); r.font.color.rgb = GRIS
    doc.add_page_break()

    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        s = line.strip()
        if not s:
            i += 1; continue
        # Tablas
        if s.startswith("|") and i+1 < len(lines) and is_sep(lines[i+1]):
            header = split_row(s)
            rows = []
            j = i + 2
            while j < len(lines) and lines[j].strip().startswith("|"):
                rows.append(split_row(lines[j])); j += 1
            make_table(doc, header, rows)
            i = j; continue
        # Encabezados
        m = re.match(r"^(#{1,4})\s+(.*)$", s)
        if m:
            lvl = min(len(m.group(1)), 3)
            txt = re.sub(r"[*`>]", "", m.group(2)).strip()
            doc.add_heading(txt, level=lvl)
            i += 1; continue
        # Blockquote
        if s.startswith(">"):
            p = doc.add_paragraph()
            add_runs(p, s.lstrip(">").strip())
            for r in p.runs:
                r.italic = True; r.font.color.rgb = GRIS; r.font.size = Pt(10)
            i += 1; continue
        # Regla horizontal
        if s in ("---", "***", "___"):
            i += 1; continue
        # Code fence
        if s.startswith("```"):
            j = i + 1; buf = []
            while j < len(lines) and not lines[j].strip().startswith("```"):
                buf.append(lines[j]); j += 1
            p = doc.add_paragraph()
            r = p.add_run("\n".join(buf)); r.font.name = "Consolas"; r.font.size = Pt(9)
            i = j + 1; continue
        # Bullets
        mb = re.match(r"^[-*]\s+(.*)$", s)
        if mb:
            p = doc.add_paragraph(style="List Bullet")
            add_runs(p, mb.group(1)); i += 1; continue
        mn = re.match(r"^\d+\.\s+(.*)$", s)
        if mn:
            p = doc.add_paragraph(style="List Number")
            add_runs(p, mn.group(1)); i += 1; continue
        # Parrafo normal
        p = doc.add_paragraph()
        add_runs(p, s)
        i += 1

    os.makedirs(os.path.dirname(dest), exist_ok=True)
    doc.save(dest)
    print("OK ->", dest)

DOCS = "/Users/jhonytoro/Documents/emprendimiento/tym_music/docs"
jobs = [
    ("/Users/jhonytoro/.claude/plans/immutable-enchanting-lemur.md",
     f"{DOCS}/TYM_Music_01_Estudio_de_Mercado.docx",
     "Estudio de Mercado y Viabilidad",
     "Benchmark · Competidores · Oceano rojo/azul · Warnings · Colombia · Junio 2026"),
    (f"{DOCS}/00-MASTER-BRIEF.md",
     f"{DOCS}/TYM_Music_00_Master_Brief.docx",
     "Master Brief — Resumen Ejecutivo de Decisiones",
     "Documento maestro del proyecto · Junio 2026"),
    (f"{DOCS}/02-modelo-negocio-y-costos.md",
     f"{DOCS}/TYM_Music_03_Modelo_y_Costos.docx",
     "Modelo de Negocio y Analisis de Costos",
     "Pay-per-priority · Unit economics · Costos operativos · Junio 2026"),
]
for src, dest, title, sub in jobs:
    if os.path.exists(src):
        convert(src, dest, title, sub)
    else:
        print("FALTA:", src)
print("Listo.")
