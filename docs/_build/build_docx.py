#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Genera el documento Word de analisis de mercado y modelo de negocio de TYM Music."""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(__file__)
CH = os.path.join(HERE, "charts")
OUT = "/Users/jhonytoro/Documents/emprendimiento/tym_music/docs/TYM_Music_02_Analisis_Completo_con_Graficas.docx"

AZUL = RGBColor(0x1F, 0x4E, 0x79)
AZUL2 = RGBColor(0x2E, 0x75, 0xB6)
GRISTXT = RGBColor(0x40, 0x40, 0x40)
BLANCO = RGBColor(0xFF, 0xFF, 0xFF)

doc = Document()
# Estilo base
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)
normal.font.color.rgb = RGBColor(0x22, 0x22, 0x22)

for lvl, color, size in [("Heading 1", AZUL, 15), ("Heading 2", AZUL2, 12.5)]:
    st = doc.styles[lvl]
    st.font.color.rgb = color
    st.font.size = Pt(size)
    st.font.name = "Calibri"

def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear"); sh.set(qn("w:color"), "auto"); sh.set(qn("w:fill"), hexcolor)
    tcPr.append(sh)

def add_table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        shade(hdr[i], "1F4E79")
        p = hdr[i].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h); r.bold = True; r.font.color.rgb = BLANCO; r.font.size = Pt(9.5)
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for i, val in enumerate(row):
            if ri % 2 == 1:
                shade(cells[i], "EAF1F8")
            p = cells[i].paragraphs[0]
            r = p.add_run(str(val)); r.font.size = Pt(9.5)
            if i == 0:
                r.bold = True
    doc.add_paragraph()
    return t

def img(name, width=6.2):
    path = os.path.join(CH, name)
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

def caption(text):
    p = doc.add_paragraph()
    r = p.add_run(text); r.italic = True; r.font.size = Pt(8.5); r.font.color.rgb = GRISTXT
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def bullets(items):
    for it in items:
        doc.add_paragraph(it, style="List Bullet")

# ---------------- PORTADA ----------------
for _ in range(3): doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("TYM MUSIC"); r.bold = True; r.font.size = Pt(40); r.font.color.rgb = AZUL
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Analisis de Mercado, Viabilidad y Modelo de Negocio"); r.font.size = Pt(16); r.font.color.rgb = AZUL2
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Sistema de musica por demanda para establecimientos comerciales — Colombia"); r.font.size = Pt(11); r.italic = True
for _ in range(2): doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Junio 2026  ·  Fase pre-MVP  ·  Documento de trabajo"); r.font.size = Pt(10); r.font.color.rgb = GRISTXT
doc.add_page_break()

# ---------------- RESUMEN EJECUTIVO ----------------
doc.add_heading("Resumen ejecutivo", level=1)
doc.add_paragraph(
    "TYM Music es un sistema que permite a los clientes de establecimientos comerciales (gimnasios, "
    "cafes, barberias, bares, etc.) poner su propia musica desde el celular mediante una web app sin "
    "descargas. El modelo de monetizacion es pago por prioridad (pay-per-priority): quien quiere que su "
    "cancion suene antes paga un micro-cobro variable que el comercio fija ($200 a $2.000 COP) y que se "
    "suma a la cuenta final del cliente. Si nadie paga, suena una lista de respaldo curada por el dueno.")
doc.add_paragraph("Conclusiones clave:")
bullets([
    "Mercado real y en crecimiento: +40.000 establecimientos nocturnos en Colombia y un subsector de bares/discotecas que crecio +64,3% en ingresos promedio en 2025.",
    "El concepto YA existe en Colombia (Melao, Apollo, Secret DJ), pero como feature secundaria y en zona gris legal. Hay espacio para un producto dedicado, legal y monetizable.",
    "Veredicto: VIABLE con diferenciacion. El oceano azul esta en 'musica legal + monetizacion para el comercio', no en clonar a Melao.",
    "Sin cobro inicial al comercio: adopta porque gana dinero sin arriesgar. La monetizacion por suscripcion llega despues, al crecer el volumen.",
    "Sumar el micro-cobro a la cuenta (en vez de cobrar por cancion) es la decision que hace rentable el modelo: baja la comision de pasarela de ~55% a <3%.",
    "Punto de equilibrio estimado: ~76 locales activos en fase embed (musica de bajo costo via reproductores embebidos).",
])

# ---------------- 1. CONCEPTO ----------------
doc.add_heading("1. Concepto y producto", level=1)
doc.add_paragraph(
    "Web app accesible por QR/enlace, sin descargas, para eliminar friccion en el primer contacto. "
    "El cliente pide canciones a una cola compartida que se reproduce en el local. A futuro, una app/ecosistema "
    "opcional para usuarios recurrentes.")
doc.add_heading("Mecanica del producto", level=2)
bullets([
    "Quien quiere escuchar, paga: micro-cobro por prioridad para adelantar su cancion.",
    "El cobro se acumula y se suma a la cuenta final del cliente (no se procesa cancion por cancion).",
    "Si no hay prioridades pagas en cola, suena la lista de respaldo curada por el dueno (legal y acorde al estilo).",
    "Control total del dueno: estilo/generos permitidos, aprobar/rechazar, volumen, horarios.",
    "Reglas anti-cola-infinita: limite por usuario, tiempo maximo de espera, precio dinamico opcional.",
])

# ---------------- 2. MERCADO ----------------
doc.add_heading("2. Mercado (Colombia)", level=1)
add_table(["Metrica", "Dato", "Fuente"], [
    ["Establecimientos nocturnos", "+40.000 (+200.000 empleos directos)", "Asobares"],
    ["Empleo bares/gastrobares/discotecas", "~96.000 empleos/mes", "DANE"],
    ["Crecimiento sector H1 2025", "+1,9% ingresos reales", "Asobares"],
    ["Subsector bares/discotecas 2025", "+13,6% ocupacion / +64,3% ingresos", "Asobares"],
    ["Informalidad en restaurantes", "~80% informales", "DANE"],
    ["Establecimientos independientes", "~95% del mercado", "ACODRES"],
])
img("01_crecimiento_sector.png")
caption("Figura 1. El subsector de bares y discotecas fue uno de los de mayor recuperacion en 2025.")
doc.add_heading("TAM / SAM / SOM (orden de magnitud)", level=2)
add_table(["Nivel", "Definicion", "Estimacion"], [
    ["TAM", "Todos los establecimientos con musica en Colombia", "~50.000+ locales; ~$48.000M COP/ano potencial"],
    ["SAM", "Formales, urbanos, con ambiente musical (4 ciudades)", "8.000–15.000 locales"],
    ["SOM", "Alcanzable ano 1–2 de forma realista", "100–400 locales"],
])
doc.add_paragraph("Riesgo de tamano: la alta informalidad (80%) reduce el mercado realmente pagador. La metrica que importa no es locales registrados sino prioridades pagadas por noche.")

# ---------------- 3. COMPETENCIA ----------------
doc.add_heading("3. Competencia y benchmark", level=1)
doc.add_heading("Competidores en Colombia (directos)", level=2)
add_table(["Producto", "Como funciona", "Notas"], [
    ["Melao", "QR en mesa, links Spotify/YouTube, cola, carta digital", "Es un POS todo-en-uno; la musica es feature secundaria. 1 bar visible. Zona gris legal."],
    ["Apollo App", "Rockola en tiempo real, votar, ver menu", "Creada por colombianos (2021). Validar si sigue activa."],
    ["'Programa tu musica'", "Pedir canciones en bares/discotecas/emisoras", "+50 establecimientos en Antioquia."],
    ["Secret DJ", "Rockola anonima desde el celular", "Presencia en Colombia."],
])
doc.add_heading("Referentes internacionales", level=2)
add_table(["Empresa", "Modelo", "Datos clave"], [
    ["TouchTunes (USA)", "Jukebox + app, pay-per-song", "~65.000 locales, 12M reproducciones/sem, 2M MAU"],
    ["Rockbot (USA)", "Musica de fondo licenciada + requests", "US$25/mes por zona; cubre licencias ASCAP/BMI/SESAC"],
    ["Soundtrack Your Brand", "Musica de fondo legal B2B", "+125M canciones licenciadas para uso comercial"],
])
img("02_posicionamiento.png", width=5.6)
caption("Figura 2. Los competidores locales operan con musica secundaria y en zona gris. TYM apunta al cuadrante de producto de musica dedicado, monetizable y legal.")

# ---------------- 4. OCEANO ----------------
doc.add_heading("4. Oceano rojo vs oceano azul", level=1)
doc.add_paragraph("Oceano ROJO (saturado): 'app/QR para pedir canciones con cola y control del dueno' — ya lo hace Melao y otros.")
doc.add_paragraph("Oceano AZUL (diferenciacion defendible):")
bullets([
    "Legalidad como producto: gestionar Sayco-Acinpro + evolucionar a catalogo licenciado.",
    "Monetizacion pay-per-priority que genera ingreso nuevo para el comercio.",
    "Producto de musica dedicado (no un add-on de un POS).",
    "Nichos desatendidos fuera del foco bar/discoteca.",
])

# ---------------- 5. MODELO DE NEGOCIO ----------------
doc.add_heading("5. Modelo de negocio", level=1)
doc.add_paragraph(
    "Monetizacion 100% por pay-per-priority, sin cobro inicial ni mensualidad al comercio en la fase de "
    "entrada. El comercio adopta porque gana dinero sin arriesgar. La suscripcion anual minima llega despues, "
    "cuando el volumen y los ingresos lo justifiquen.")
doc.add_heading("Por que 'sumar a la cuenta' es la decision correcta", level=2)
add_table(["Concepto", "Micropago por cancion ($1.500)", "Acumulado a la cuenta"], [
    ["Comision pasarela (Wompi tarjeta)", "2,65% + $700 + IVA ≈ $896", "2,65% sobre el total mensual"],
    ["% que se come la comision", "~45–60%", "<3%"],
    ["Friccion para el cliente", "Pagar cada vez", "Cero: paga todo junto al final"],
])
img("03_comision_micropago.png")
caption("Figura 3. Procesar cada cancion como micropago es inviable; acumular y liquidar en bloque hace rentable el modelo.")
doc.add_heading("Flujo de dinero (settlement)", level=2)
bullets([
    "Cliente paga su cuenta completa al comercio (incluye el $ de prioridades).",
    "Comercio liquida en bloque (semanal/mensual) la parte de TYM.",
    "Recomendado: el comercio retiene su % y paga solo lo de TYM (menos friccion, mas confianza).",
    "Riesgo de cartera: mitigar con liquidacion frecuente, garantia pequena o corte por mora.",
])

# ---------------- 6. PRECIO ----------------
doc.add_heading("6. Precio variable por comercio", level=1)
doc.add_paragraph(
    "El precio por prioridad NO es fijo: cada comercio lo define segun su publico y ticket, desde $200 COP "
    "(cafes, salas de espera) hasta $2.000 COP (locales premium / discotecas). El costo para el usuario debe "
    "sentirse muy bajo porque va sumado a su cuenta.")
img("04_precios_tiers.png", width=5.8)
caption("Figura 4. Rango de precio por prioridad segun tipo de comercio.")

# ---------------- 7. LEGAL ----------------
doc.add_heading("7. Estrategia legal y de fuentes de musica", level=1)
doc.add_paragraph(
    "Principio rector: desde el principio NO meterse en problemas legales. Estrategia por fases:")
add_table(["Fase", "Fuente de musica", "Enfoque legal"], [
    ["Entrada (MVP)", "Reproductores oficiales embebidos (YouTube/Spotify) con front desacoplado del back",
     "No redistribuimos audio ni usamos APIs prohibidas; el reproductor oficial corre del lado del cliente. Se gestiona/recomienda licencia Sayco-Acinpro del local."],
    ["Crecimiento", "Lista de respaldo desde catalogo licenciado", "Reduce dependencia de fuentes externas; mayoria del tiempo suena musica con derechos."],
    ["Madurez", "Catalogo licenciado propio o revendido para uso comercial", "Elimina la zona gris; diferenciador defendible (modelo Rockbot/Soundtrack)."],
])
doc.add_paragraph(
    "Importante: en Colombia toda comunicacion publica de musica requiere licencia de la Organizacion Sayco y "
    "Acinpro (OSA), con tarifa variable y negociable. TYM puede gestionarla como servicio integrado. "
    "La estrategia de fuente debe validarse con un abogado de propiedad intelectual antes de escalar.")

# ---------------- 8. COSTOS ----------------
doc.add_heading("8. Costos operativos", level=1)
add_table(["Tipo", "Costo", "Notas"], [
    ["Fijo", "Equipo / desarrollo", "Costo dominante en etapa inicial"],
    ["Fijo", "Infraestructura cloud (tiempo real, DB)", "US$50–250/mes en MVP; escala con uso"],
    ["Variable", "Fuente de musica", "$0 en fase embed; US$15–30/local/mes con catalogo licenciado"],
    ["Variable", "Pasarela de pago", "~2,65% + $700 por liquidacion mensual por local (despreciable por cancion)"],
    ["Variable", "Licencia Sayco-Acinpro", "Variable; la paga el comercio, TYM puede gestionarla"],
])
img("09_costos.png", width=5.0)
caption("Figura 5. Estructura de costos operativos (ilustrativa). El equipo/desarrollo domina al inicio; el catalogo licenciado entra en fase 2.")
doc.add_paragraph(
    "La fuente de musica es la variable que mas mueve el margen: empezar con embeds (costo ~$0) permite "
    "alcanzar el punto de equilibrio antes y financiar la migracion a catalogo licenciado.")

# ---------------- 9. UNIT ECONOMICS ----------------
doc.add_heading("9. Unit economics", level=1)
doc.add_paragraph("Supuestos: precio promedio ponderado $700/prioridad, ~25 prioridades/noche, ~18 noches/mes "
                  "(450 prioridades/mes), split comercio 35% / TYM 65%.")
add_table(["Concepto", "Por local activo / mes"], [
    ["Recaudo bruto (450 x $700)", "$315.000"],
    ["Parte del comercio (35%)", "$110.250"],
    ["Ingreso bruto TYM (65%)", "$204.750"],
    ["– Pasarela (2,65%)", "~$8.350"],
    ["– Musica (fase embed)", "$0"],
    ["Margen de contribucion TYM (embed)", "~$196.400"],
    ["Margen de contribucion TYM (catalogo licenciado)", "~$116.400"],
])
img("05_unit_economics.png", width=5.8)
caption("Figura 6. Descomposicion del recaudo por local activo en la fase embed.")

# ---------------- 10. PROYECCIONES ----------------
doc.add_heading("10. Proyecciones y punto de equilibrio", level=1)
img("06_proyeccion_escala.png", width=5.8)
caption("Figura 7. Margen de contribucion mensual segun numero de locales activos y fuente de musica.")
img("07_break_even.png", width=5.8)
caption("Figura 8. Punto de equilibrio: ~76 locales activos en fase embed; ~129 con catalogo licenciado "
        "(asumiendo costos fijos de ~$15M COP/mes).")
doc.add_paragraph("La metrica critica no es locales registrados, sino locales ACTIVOS (con prioridades pagadas por noche). "
                  "Empezar con embeds reduce el punto de equilibrio casi a la mitad.")

# ---------------- 11. ROADMAP MONETIZACION ----------------
doc.add_heading("11. Roadmap de monetizacion", level=1)
bullets([
    "Ano 1 — Entrar: 100% pay-per-priority, sin cobro al comercio. Objetivo: traccion y locales activos.",
    "Ano 2 — Crecer: introducir suscripcion anual minima que cubra gastos operativos, una vez probado el valor.",
    "Ano 3 — Consolidar: catalogo licenciado, ecosistema de usuario (app), expansion de nichos.",
])
img("08_roadmap_monetizacion.png", width=5.8)
caption("Figura 9. Evolucion de ingresos: entrar gratis para el comercio y agregar suscripcion al crecer (cifras ilustrativas).")

# ---------------- 12. NICHOS ----------------
doc.add_heading("12. Estrategia de nichos", level=1)
doc.add_paragraph("Estrategia: atacar varios nichos gradualmente y dejar que los datos muestren donde tracciona mas, "
                  "en lugar de competir de frente con Melao en bares.")
add_table(["Nicho", "Por que encaja"], [
    ["Gimnasios / crossfit", "Musica central, publico joven, alto engagement con 'que suene mi cancion'"],
    ["Cafes / coworkings", "Ambiente importa; publico dispuesto a micropagar; diurno y recurrente"],
    ["Peluquerias / barberias / spa", "Sesiones largas = tiempo muerto monetizable; sin competencia"],
    ["Food courts / cervecerias", "Mucha gente, rotacion, volumen"],
    ["Salas de espera / lavanderias", "Tiempo muerto, cero competencia"],
])

# ---------------- 13. RIESGOS ----------------
doc.add_heading("13. Riesgos y mitigaciones", level=1)
add_table(["Riesgo", "Mitigacion"], [
    ["Legal de fuentes (Spotify/YouTube prohiben uso comercial)", "Embeds de reproductores oficiales + migracion a catalogo licenciado + gestion Sayco-Acinpro"],
    ["Cartera: el comercio no liquida", "Liquidacion frecuente, garantia pequena, corte por mora"],
    ["Pocos pagan (ingreso depende de prioridades)", "Fallback random mantiene el servicio; validar willingness-to-pay en campo"],
    ["Disputas en la cuenta", "Confirmacion clara al pedir, registro, limite por usuario"],
    ["Concepto ya existe (Melao, etc.)", "Diferenciar por musica dedicada + legalidad + ingreso para el comercio"],
    ["Alta informalidad del mercado", "Foco en locales activos y nichos; modelo sin costo fijo para el comercio"],
])

# ---------------- 14. SUPUESTOS ----------------
doc.add_heading("14. Supuestos del analisis", level=1)
bullets([
    "Precio promedio ponderado por prioridad: $700 COP (rango real $200–$2.000 segun comercio).",
    "Local activo: ~25 prioridades/noche x ~18 noches/mes = 450 prioridades/mes.",
    "Split comercio 35% / TYM 65% (palanca a validar).",
    "Costos fijos de equipo pequeno: ~$15M COP/mes (estimado a validar).",
    "Comision pasarela: 2,65% + $700 + IVA (Wompi, tarjeta; PSE 1,49%).",
    "Costo catalogo licenciado: US$15–30/local/mes (a confirmar con proveedores).",
    "Las cifras de proyeccion de ingresos por ano son ilustrativas, no proyeccion financiera final.",
])

# ---------------- 15. PROXIMOS PASOS ----------------
doc.add_heading("15. Proximos pasos", level=1)
bullets([
    "Validacion de campo: 10–15 entrevistas a duenos de locales en varios nichos (guion ya definido).",
    "Definir con abogado PI el 'sistema intermedio' legal de fuentes (embeds desacoplados).",
    "Confirmar split, precio promedio y costos reales (cloud, catalogo, pasarela).",
    "Construir MVP web app: cola en tiempo real + panel del dueno + QR cliente + registro de cobros a la cuenta.",
    "Piloto en 1–2 nichos para medir prioridades pagadas por noche (la metrica clave).",
])

# ---------------- FUENTES ----------------
doc.add_heading("Fuentes", level=1)
for s in [
    "TouchTunes / Rockbot / BarBox: business.touchtunes.com, rockbot.com, tracxn.com, pitchbook.com",
    "Colombia: melao.app, infobae (Apollo App), enter.co (Secret DJ), telemedellin.tv",
    "Legal: support.spotify.com, soundtrack.io, osa.org.co, sayco.org, acinpro.org.co, saycoacinpro.org",
    "Pasarelas: wompi.com, soporte.wompi.co, guiadesoftware.com",
    "Catalogo legal: soundtrack.io, jamendo licensing, epidemicsound.com",
    "Mercado: asobares.org, eltiempo.com, revistalabarra.com, acoga.org, semana.com, DANE/ACODRES",
]:
    p = doc.add_paragraph(s, style="List Bullet"); p.runs[0].font.size = Pt(9)

os.makedirs(os.path.dirname(OUT), exist_ok=True)
doc.save(OUT)
print("OK - documento guardado en", OUT)
